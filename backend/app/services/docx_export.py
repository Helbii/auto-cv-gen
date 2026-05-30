"""
Génération d'un CV Word (.docx) à partir des données structurées.
Produit un document éditable, prêt pour Word / LibreOffice.
"""
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .audit import resolve_safe_summary, resolve_safe_title
from .markdown import (
    _format_achievement,
    _format_cv_date,
    _format_location,
    _is_database_cv,
    group_experience_bullets,
    stack_for_source,
    build_offer_terms,
    rank_missions_by_relevance,
    select_relevant_projects,
    format_project_entry,
)

# ── Palette ───────────────────────────────────────────────────────────────────
BLUE    = RGBColor(0,   79, 144)
GRAY    = RGBColor(100, 100, 100)
BLACK   = RGBColor(0,   0,   0)
FONT    = "Calibri"
BODY_PT = 10.0


# ── Helpers bas niveau ────────────────────────────────────────────────────────

def _run(para, text: str, size: float = BODY_PT, bold: bool = False,
         color: RGBColor = BLACK, italic: bool = False):
    r = para.add_run(text)
    r.font.name  = FONT
    r.font.size  = Pt(size)
    r.font.bold  = bold
    r.font.italic = italic
    r.font.color.rgb = color
    return r


def _para(doc: Document, space_before: float = 0, space_after: float = 3) -> Any:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(BODY_PT * 1.25)
    return p


def _bottom_border(para, color_hex: str = "004F90", thickness: str = "6") -> None:
    """Ajoute un filet bleu sous un paragraphe (style section header)."""
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    thickness)
    bot.set(qn("w:space"), "2")
    bot.set(qn("w:color"), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)


def _section_title(doc: Document, text: str) -> Any:
    p = _para(doc, space_before=10, space_after=4)
    _run(p, text.upper(), size=10.5, bold=True, color=BLUE)
    _bottom_border(p)
    return p


def _bullet(doc: Document, text: str) -> Any:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.5)
    _run(p, text, size=BODY_PT)
    return p


def _set_page_margins(doc: Document, margin_cm: float = 1.8) -> None:
    for section in doc.sections:
        m = Cm(margin_cm)
        section.top_margin    = m
        section.bottom_margin = m
        section.left_margin   = m
        section.right_margin  = m


# ── Sections contenu ──────────────────────────────────────────────────────────

def _add_header_database(doc: Document, cv_master: Dict[str, Any],
                          matching: Dict[str, Any], generated: Dict[str, Any],
                          audit: Dict[str, Any]) -> None:
    profile  = cv_master.get("profile", {})
    contact  = profile.get("contact", {})
    full_name = " ".join(p for p in [profile.get("firstName"), profile.get("lastName")] if p)
    title     = resolve_safe_title(matching, generated, audit)

    # Nom
    p_name = _para(doc, space_before=0, space_after=2)
    _run(p_name, full_name, size=20, bold=True, color=BLUE)

    # Titre du poste
    p_title = _para(doc, space_before=0, space_after=4)
    _run(p_title, title, size=11, italic=True, color=GRAY)

    # Contacts sur une ligne
    contact_parts = [
        _format_location(profile.get("location")),
        contact.get("email"),
        contact.get("phone"),
        contact.get("linkedin"),
        contact.get("github"),
    ]
    contact_str = "  |  ".join(p for p in contact_parts if p)
    if contact_str:
        p_contact = _para(doc, space_before=0, space_after=6)
        _run(p_contact, contact_str, size=8.5, color=GRAY)


def _add_summary(doc: Document, matching: Dict[str, Any],
                 generated: Dict[str, Any], audit: Dict[str, Any]) -> None:
    summary = resolve_safe_summary(generated, audit, matching)
    if not summary:
        return
    _section_title(doc, "Résumé professionnel")
    p = _para(doc, space_after=4)
    _run(p, summary, size=BODY_PT)


def _add_skills_database(doc: Document, cv_master: Dict[str, Any],
                          audit: Dict[str, Any]) -> None:
    technical = cv_master.get("skills", {}).get("technical", {}) if isinstance(cv_master.get("skills"), dict) else {}
    skill_groups = [
        ("Compétences ciblées",  None),
        ("Langages",             "programmingLanguages"),
        ("Backend & APIs",       "backend"),
        ("Data & IA",            "dataEngineering"),
        ("Bases de données",     "databases"),
        ("Frontend",             "frontend"),
        ("DevOps & outils",      "devOpsTools"),
        ("IoT & embarqué",       "embeddedIoT"),
    ]

    has_content = False
    lines: List[str] = []

    validated = [it.get("skill", "") for it in audit.get("valid_skills", []) if it.get("skill")]
    if validated:
        lines.append(f"Compétences ciblées : {', '.join(dict.fromkeys(validated[:10]))}")
        has_content = True

    for label, key in skill_groups[1:]:
        values = technical.get(key, []) if isinstance(technical, dict) and key else []
        if values:
            lines.append(f"{label} : {', '.join(values[:8])}")
            has_content = True

    if not has_content:
        return

    _section_title(doc, "Compétences clés")
    for line in lines:
        _bullet(doc, line)


def _add_experiences_database(doc: Document, cv_master: Dict[str, Any], matching: Optional[Dict[str, Any]] = None) -> None:
    experiences = cv_master.get("experiences", [])
    if not experiences:
        return
    _section_title(doc, "Expériences professionnelles")
    offer_terms = build_offer_terms(matching or {})
    for index, exp in enumerate(experiences[:3]):
        date     = _format_cv_date(exp.get("startDate"), exp.get("endDate"), bool(exp.get("isCurrent")))
        location = _format_location(exp.get("location"))
        header_parts = [exp.get("position"), exp.get("company"), location, date]
        header_str   = " — ".join(p for p in header_parts if p)

        p_head = _para(doc, space_before=6, space_after=2)
        _run(p_head, header_str, size=10.5, bold=True)

        if index == 0 and exp.get("summary"):
            p_sum = _para(doc, space_after=2)
            _run(p_sum, exp["summary"], size=BODY_PT, italic=True)

        # Missions triées par pertinence, volume réduit pour laisser place aux projets
        mission_limit = 4 if index == 0 else 3
        for mission in rank_missions_by_relevance(exp.get("missions", []), offer_terms, mission_limit):
            _bullet(doc, mission)

        for achievement in exp.get("achievements", [])[:(1 if index <= 1 else 0)]:
            text = _format_achievement(achievement)
            if text:
                _bullet(doc, text)

        technologies = exp.get("technologies", [])
        if technologies:
            _bullet(doc, f"Stack : {', '.join(technologies[:10])}")


def _add_projects_database(doc: Document, cv_master: Dict[str, Any], matching: Dict[str, Any]) -> None:
    projects = select_relevant_projects(cv_master, matching, max_projects=3)
    if not projects:
        return
    _section_title(doc, "Projets")
    for proj in projects:
        info = format_project_entry(proj)
        header = info["name"]
        if info["date"]:
            header = f"{header} — {info['date']}"
        p_head = _para(doc, space_before=5, space_after=2)
        _run(p_head, header, size=10.5, bold=True)
        if info["description"]:
            _bullet(doc, info["description"])
        if info["technologies"]:
            _bullet(doc, f"Stack : {', '.join(info['technologies'])}")


def _add_experiences_evidence(
    doc: Document,
    audit: Dict[str, Any],
    id_to_evidence: Dict[str, Any],
) -> None:
    grouped = group_experience_bullets(audit.get("valid_bullets", []))
    if not grouped:
        return
    _section_title(doc, "Expériences professionnelles")
    for experience in grouped:
        p_head = _para(doc, space_before=6, space_after=2)
        _run(p_head, experience["source"], size=10.5, bold=True)
        for bullet in experience["bullets"][:6]:
            _bullet(doc, bullet)
        stack = stack_for_source(experience["source"], id_to_evidence)
        if stack:
            _bullet(doc, f"Stack : {stack}")


def _add_education_database(doc: Document, cv_master: Dict[str, Any]) -> None:
    education = cv_master.get("education", [])
    if not education:
        return
    _section_title(doc, "Formation")
    for edu in education[:4]:
        date  = _format_cv_date(edu.get("startDate"), edu.get("endDate"), bool(edu.get("isCurrent")))
        parts = [edu.get("school"), edu.get("degree"), date]
        _bullet(doc, " — ".join(p for p in parts if p))


def _add_languages_database(doc: Document, cv_master: Dict[str, Any]) -> None:
    languages = []
    for lang in cv_master.get("languages", []):
        if lang.get("language") and lang.get("level"):
            languages.append(f"{lang['language']} : {lang['level']}")
    certs = [c.get("name") for c in cv_master.get("certifications", []) if c.get("name")]
    if certs and len(languages) >= 2:
        languages[1] = f"{languages[1]} — {certs[0]}"
    if not languages:
        return
    _section_title(doc, "Langues")
    for lang in languages:
        _bullet(doc, lang)


# ── Point d'entrée ────────────────────────────────────────────────────────────

def generate_cv_docx(
    matching: Dict[str, Any],
    generated: Dict[str, Any],
    audit: Dict[str, Any],
    id_to_evidence: Dict[str, Any],
    cv_master: Optional[Dict[str, Any]],
    output_path: Path,
) -> Path:
    doc = Document()
    _set_page_margins(doc, margin_cm=1.8)

    # Supprime le style de paragraphe par défaut (espacement excessif)
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(BODY_PT)
    style.paragraph_format.space_after  = Pt(3)
    style.paragraph_format.space_before = Pt(0)

    if _is_database_cv(cv_master):
        _add_header_database(doc, cv_master, matching, generated, audit)
        _add_summary(doc, matching, generated, audit)
        _add_skills_database(doc, cv_master, audit)
        _add_experiences_database(doc, cv_master, matching)
        _add_projects_database(doc, cv_master, matching)
        _add_education_database(doc, cv_master)
        _add_languages_database(doc, cv_master)
    else:
        title = resolve_safe_title(matching, generated, audit)
        p_title = _para(doc, space_before=0, space_after=6)
        _run(p_title, title, size=18, bold=True, color=BLUE)

        _add_summary(doc, matching, generated, audit)

        skills = [it.get("skill", "") for it in audit.get("valid_skills", []) if it.get("skill")]
        if skills:
            _section_title(doc, "Compétences clés")
            for skill in skills:
                _bullet(doc, skill)

        _add_experiences_evidence(doc, audit, id_to_evidence)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
